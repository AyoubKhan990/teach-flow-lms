import sys
import json
import os
import re
from docx import Document

def generate_docx(data):
    doc = Document()
    
    # Title
    title = data.get('topic', 'Assignment')
    heading = doc.add_heading(title, 0)
    heading.alignment = 1 # Center

    # Metadata
    if 'subject' in data:
        p = doc.add_paragraph()
        runner = p.add_run(f"Subject: {data['subject']}")
        runner.bold = True
    if 'level' in data:
        p = doc.add_paragraph()
        runner = p.add_run(f"Level: {data['level']}")
        runner.bold = True
    
    doc.add_paragraph() # Spacer

    # Content
    content = data.get('content', '')

    lines = content.split('\n')

    def add_inline_runs(paragraph, text):
        bold_parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in bold_parts:
            if part.startswith('**') and part.endswith('**') and len(part) >= 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                continue

            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and len(ip) >= 2 and not ip.startswith('**'):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ip)

    # Parse and add content
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\[IMAGE:\s*(.*?)\]$', line):
            continue

        # Parse Markdown Headings
        if line.startswith('#### '):
            doc.add_heading(line.replace('#### ', ''), level=4)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        # Parse Lists
        elif line.startswith('- '):
            p = doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        elif re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s+', '', line), style='List Number')
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)

    # Save to file
    output_filename = f"generated_{data.get('id', 'temp')}.docx"
    doc.save(output_filename)
    print(output_filename)

if __name__ == "__main__":
    try:
        input_data = sys.stdin.read()
        data = json.loads(input_data)
        generate_docx(data)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
